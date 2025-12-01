"""
Document Processing for RAG System
Supports multiple file formats with intelligent chunking strategies
"""
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass
from enum import Enum
from pathlib import Path
import logging
import hashlib
import re
from datetime import datetime

logger = logging.getLogger(__name__)


class ChunkStrategy(Enum):
    """Chunking strategies for document processing"""
    FIXED_SIZE = "fixed_size"
    SEMANTIC = "semantic"
    SENTENCE = "sentence"
    PARAGRAPH = "paragraph"
    SLIDING_WINDOW = "sliding_window"


class FileType(Enum):
    """Supported file types"""
    TXT = "txt"
    PDF = "pdf"
    DOCX = "docx"
    MD = "md"
    MARKDOWN = "markdown"
    JSON = "json"
    CSV = "csv"
    HTML = "html"
    UNKNOWN = "unknown"


@dataclass
class Document:
    """Represents a processed document"""
    content: str
    metadata: Dict[str, Any]
    file_path: Optional[str] = None
    file_type: Optional[FileType] = None
    hash: Optional[str] = None

    def __post_init__(self):
        """Generate hash if not provided"""
        if not self.hash:
            self.hash = hashlib.sha256(self.content.encode()).hexdigest()


@dataclass
class DocumentChunk:
    """Represents a chunk of a document"""
    content: str
    metadata: Dict[str, Any]
    chunk_id: int
    document_hash: str
    start_char: int
    end_char: int


@dataclass
class ChunkConfig:
    """Configuration for chunking"""
    strategy: ChunkStrategy = ChunkStrategy.FIXED_SIZE
    chunk_size: int = 512
    chunk_overlap: int = 50
    min_chunk_size: int = 100
    max_chunk_size: int = 2000
    respect_sentences: bool = True
    respect_paragraphs: bool = True


class FileReader:
    """Reads files of different formats"""

    @staticmethod
    def read_txt(file_path: Path) -> str:
        """Read plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()

    @staticmethod
    def read_pdf(file_path: Path) -> str:
        """Read PDF file"""
        try:
            import PyPDF2
            text = []
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page in pdf_reader.pages:
                    text.append(page.extract_text())
            return "\n\n".join(text)
        except ImportError:
            logger.error("PyPDF2 not installed. Install with: pip install PyPDF2")
            raise
        except Exception as e:
            logger.error(f"Error reading PDF: {e}")
            raise

    @staticmethod
    def read_docx(file_path: Path) -> str:
        """Read DOCX file"""
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            paragraphs = [para.text for para in doc.paragraphs]
            return "\n\n".join(paragraphs)
        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            raise
        except Exception as e:
            logger.error(f"Error reading DOCX: {e}")
            raise

    @staticmethod
    def read_markdown(file_path: Path) -> str:
        """Read Markdown file"""
        return FileReader.read_txt(file_path)

    @staticmethod
    def read_json(file_path: Path) -> str:
        """Read JSON file"""
        import json
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        # Convert JSON to readable text
        if isinstance(data, dict):
            # Format dict as key-value pairs
            lines = []
            for key, value in data.items():
                lines.append(f"{key}: {value}")
            return "\n".join(lines)
        elif isinstance(data, list):
            # Format list items
            return "\n\n".join(str(item) for item in data)
        else:
            return str(data)

    @staticmethod
    def read_csv(file_path: Path) -> str:
        """Read CSV file"""
        import csv
        rows = []
        with open(file_path, 'r', encoding='utf-8') as f:
            reader = csv.reader(f)
            headers = next(reader, None)

            if headers:
                rows.append(" | ".join(headers))
                rows.append("-" * 50)

            for row in reader:
                rows.append(" | ".join(row))

        return "\n".join(rows)

    @staticmethod
    def read_html(file_path: Path) -> str:
        """Read HTML file and extract text"""
        try:
            from bs4 import BeautifulSoup
            with open(file_path, 'r', encoding='utf-8') as f:
                soup = BeautifulSoup(f, 'html.parser')

            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()

            # Get text
            text = soup.get_text()

            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)

            return text
        except ImportError:
            logger.error("beautifulsoup4 not installed. Install with: pip install beautifulsoup4")
            raise
        except Exception as e:
            logger.error(f"Error reading HTML: {e}")
            raise

    @staticmethod
    def detect_file_type(file_path: Path) -> FileType:
        """Detect file type from extension"""
        suffix = file_path.suffix.lower().lstrip('.')

        type_map = {
            'txt': FileType.TXT,
            'pdf': FileType.PDF,
            'docx': FileType.DOCX,
            'doc': FileType.DOCX,
            'md': FileType.MD,
            'markdown': FileType.MARKDOWN,
            'json': FileType.JSON,
            'csv': FileType.CSV,
            'html': FileType.HTML,
            'htm': FileType.HTML
        }

        return type_map.get(suffix, FileType.UNKNOWN)

    @staticmethod
    def read_file(file_path: Path) -> Tuple[str, FileType]:
        """Read file and return content with file type"""
        file_type = FileReader.detect_file_type(file_path)

        readers = {
            FileType.TXT: FileReader.read_txt,
            FileType.PDF: FileReader.read_pdf,
            FileType.DOCX: FileReader.read_docx,
            FileType.MD: FileReader.read_markdown,
            FileType.MARKDOWN: FileReader.read_markdown,
            FileType.JSON: FileReader.read_json,
            FileType.CSV: FileReader.read_csv,
            FileType.HTML: FileReader.read_html
        }

        reader = readers.get(file_type)
        if not reader:
            raise ValueError(f"Unsupported file type: {file_type}")

        content = reader(file_path)
        return content, file_type


class TextChunker:
    """Chunks text using various strategies"""

    def __init__(self, config: ChunkConfig):
        self.config = config

    def chunk(self, text: str, document: Document) -> List[DocumentChunk]:
        """Chunk text based on configured strategy"""
        strategy_map = {
            ChunkStrategy.FIXED_SIZE: self._chunk_fixed_size,
            ChunkStrategy.SEMANTIC: self._chunk_semantic,
            ChunkStrategy.SENTENCE: self._chunk_sentence,
            ChunkStrategy.PARAGRAPH: self._chunk_paragraph,
            ChunkStrategy.SLIDING_WINDOW: self._chunk_sliding_window
        }

        chunker = strategy_map.get(self.config.strategy, self._chunk_fixed_size)
        return chunker(text, document)

    def _chunk_fixed_size(self, text: str, document: Document) -> List[DocumentChunk]:
        """Chunk text into fixed-size pieces with overlap"""
        chunks = []
        chunk_id = 0
        start = 0

        while start < len(text):
            end = start + self.config.chunk_size

            # If respecting sentences, try to end at sentence boundary
            if self.config.respect_sentences and end < len(text):
                # Look for sentence boundaries near the end
                search_start = max(start, end - 100)
                sentence_boundaries = [m.end() for m in re.finditer(r'[.!?]\s+', text[search_start:end])]
                if sentence_boundaries:
                    end = search_start + sentence_boundaries[-1]

            chunk_text = text[start:end].strip()

            if len(chunk_text) >= self.config.min_chunk_size:
                chunk = DocumentChunk(
                    content=chunk_text,
                    metadata={
                        **document.metadata,
                        "chunk_strategy": self.config.strategy.value,
                        "start_char": start,
                        "end_char": end
                    },
                    chunk_id=chunk_id,
                    document_hash=document.hash,
                    start_char=start,
                    end_char=end
                )
                chunks.append(chunk)
                chunk_id += 1

            # Move start position with overlap
            start = end - self.config.chunk_overlap

            # Ensure we're making progress
            if start >= end:
                start = end

        return chunks

    def _chunk_semantic(self, text: str, document: Document) -> List[DocumentChunk]:
        """Chunk text based on semantic boundaries (paragraphs + sentences)"""
        # Split by paragraphs first
        paragraphs = re.split(r'\n\s*\n', text)

        chunks = []
        chunk_id = 0
        current_chunk = []
        current_size = 0
        start_char = 0

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            para_size = len(para)

            # If adding this paragraph exceeds max size, create chunk
            if current_size + para_size > self.config.max_chunk_size and current_chunk:
                chunk_text = "\n\n".join(current_chunk)
                chunk = DocumentChunk(
                    content=chunk_text,
                    metadata={
                        **document.metadata,
                        "chunk_strategy": "semantic"
                    },
                    chunk_id=chunk_id,
                    document_hash=document.hash,
                    start_char=start_char,
                    end_char=start_char + len(chunk_text)
                )
                chunks.append(chunk)
                chunk_id += 1

                start_char += len(chunk_text) + 2  # +2 for newlines
                current_chunk = []
                current_size = 0

            current_chunk.append(para)
            current_size += para_size

        # Add remaining chunk
        if current_chunk:
            chunk_text = "\n\n".join(current_chunk)
            chunk = DocumentChunk(
                content=chunk_text,
                metadata={
                    **document.metadata,
                    "chunk_strategy": "semantic"
                },
                chunk_id=chunk_id,
                document_hash=document.hash,
                start_char=start_char,
                end_char=start_char + len(chunk_text)
            )
            chunks.append(chunk)

        return chunks

    def _chunk_sentence(self, text: str, document: Document) -> List[DocumentChunk]:
        """Chunk text by sentences"""
        # Split into sentences
        sentences = re.split(r'(?<=[.!?])\s+', text)

        chunks = []
        chunk_id = 0
        current_chunk = []
        current_size = 0
        start_char = 0

        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue

            sentence_size = len(sentence)

            # Check if adding this sentence exceeds chunk size
            if current_size + sentence_size > self.config.chunk_size and current_chunk:
                chunk_text = " ".join(current_chunk)
                chunk = DocumentChunk(
                    content=chunk_text,
                    metadata={
                        **document.metadata,
                        "chunk_strategy": "sentence",
                        "sentence_count": len(current_chunk)
                    },
                    chunk_id=chunk_id,
                    document_hash=document.hash,
                    start_char=start_char,
                    end_char=start_char + len(chunk_text)
                )
                chunks.append(chunk)
                chunk_id += 1

                start_char += len(chunk_text) + 1
                current_chunk = []
                current_size = 0

            current_chunk.append(sentence)
            current_size += sentence_size

        # Add remaining chunk
        if current_chunk:
            chunk_text = " ".join(current_chunk)
            chunk = DocumentChunk(
                content=chunk_text,
                metadata={
                    **document.metadata,
                    "chunk_strategy": "sentence",
                    "sentence_count": len(current_chunk)
                },
                chunk_id=chunk_id,
                document_hash=document.hash,
                start_char=start_char,
                end_char=start_char + len(chunk_text)
            )
            chunks.append(chunk)

        return chunks

    def _chunk_paragraph(self, text: str, document: Document) -> List[DocumentChunk]:
        """Chunk text by paragraphs"""
        paragraphs = re.split(r'\n\s*\n', text)

        chunks = []
        chunk_id = 0
        start_char = 0

        for para in paragraphs:
            para = para.strip()
            if not para or len(para) < self.config.min_chunk_size:
                start_char += len(para) + 2
                continue

            # If paragraph is too large, split it
            if len(para) > self.config.max_chunk_size:
                # Fall back to fixed size chunking for this paragraph
                sub_chunks = self._chunk_fixed_size(para, document)
                for sub_chunk in sub_chunks:
                    sub_chunk.chunk_id = chunk_id
                    sub_chunk.start_char += start_char
                    sub_chunk.end_char += start_char
                    chunks.append(sub_chunk)
                    chunk_id += 1
            else:
                chunk = DocumentChunk(
                    content=para,
                    metadata={
                        **document.metadata,
                        "chunk_strategy": "paragraph"
                    },
                    chunk_id=chunk_id,
                    document_hash=document.hash,
                    start_char=start_char,
                    end_char=start_char + len(para)
                )
                chunks.append(chunk)
                chunk_id += 1

            start_char += len(para) + 2

        return chunks

    def _chunk_sliding_window(self, text: str, document: Document) -> List[DocumentChunk]:
        """Chunk text with sliding window approach"""
        # Similar to fixed size but with smaller stride
        stride = self.config.chunk_size - self.config.chunk_overlap

        chunks = []
        chunk_id = 0

        for start in range(0, len(text), stride):
            end = start + self.config.chunk_size
            chunk_text = text[start:end].strip()

            if len(chunk_text) >= self.config.min_chunk_size:
                chunk = DocumentChunk(
                    content=chunk_text,
                    metadata={
                        **document.metadata,
                        "chunk_strategy": "sliding_window",
                        "stride": stride
                    },
                    chunk_id=chunk_id,
                    document_hash=document.hash,
                    start_char=start,
                    end_char=end
                )
                chunks.append(chunk)
                chunk_id += 1

        return chunks


class DocumentProcessor:
    """Main document processor for RAG system"""

    def __init__(self, chunk_config: Optional[ChunkConfig] = None):
        self.chunk_config = chunk_config or ChunkConfig()
        self.chunker = TextChunker(self.chunk_config)
        self.file_reader = FileReader()

    def process_file(self, file_path: str) -> Tuple[Document, List[DocumentChunk]]:
        """
        Process a file and return document with chunks.

        Args:
            file_path: Path to file

        Returns:
            Tuple of (Document, List of DocumentChunks)
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        # Read file
        content, file_type = self.file_reader.read_file(path)

        # Extract metadata
        metadata = self._extract_metadata(path, file_type, content)

        # Create document
        document = Document(
            content=content,
            metadata=metadata,
            file_path=str(path),
            file_type=file_type
        )

        # Chunk document
        chunks = self.chunker.chunk(content, document)

        return document, chunks

    def process_directory(
        self,
        directory_path: str,
        recursive: bool = True,
        file_patterns: Optional[List[str]] = None
    ) -> List[Tuple[Document, List[DocumentChunk]]]:
        """
        Process all files in a directory.

        Args:
            directory_path: Path to directory
            recursive: Whether to process subdirectories
            file_patterns: Optional list of glob patterns to match

        Returns:
            List of (Document, chunks) tuples
        """
        path = Path(directory_path)

        if not path.exists() or not path.is_dir():
            raise ValueError(f"Invalid directory: {directory_path}")

        results = []

        # Default patterns if none provided
        if not file_patterns:
            file_patterns = ['*.txt', '*.pdf', '*.docx', '*.md', '*.json', '*.csv', '*.html']

        # Find all matching files
        files = []
        for pattern in file_patterns:
            if recursive:
                files.extend(path.rglob(pattern))
            else:
                files.extend(path.glob(pattern))

        # Process each file
        for file_path in files:
            try:
                doc, chunks = self.process_file(str(file_path))
                results.append((doc, chunks))
                logger.info(f"Processed {file_path.name}: {len(chunks)} chunks")
            except Exception as e:
                logger.error(f"Error processing {file_path}: {e}")
                continue

        return results

    def _extract_metadata(self, path: Path, file_type: FileType, content: str) -> Dict[str, Any]:
        """Extract metadata from file"""
        stat = path.stat()

        metadata = {
            "file_name": path.name,
            "file_path": str(path),
            "file_type": file_type.value,
            "file_size": stat.st_size,
            "created_at": datetime.fromtimestamp(stat.st_ctime).isoformat(),
            "modified_at": datetime.fromtimestamp(stat.st_mtime).isoformat(),
            "content_length": len(content),
            "word_count": len(content.split()),
            "line_count": content.count('\n') + 1
        }

        # Add content-based metadata
        metadata.update(self._analyze_content(content))

        return metadata

    def _analyze_content(self, content: str) -> Dict[str, Any]:
        """Analyze content for additional metadata"""
        return {
            "has_code": bool(re.search(r'```|def |class |function |import ', content)),
            "has_urls": bool(re.search(r'https?://', content)),
            "has_emails": bool(re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', content)),
            "has_numbers": bool(re.search(r'\d+', content)),
            "language": self._detect_language(content)
        }

    def _detect_language(self, content: str) -> str:
        """Simple language detection"""
        # This is a very basic implementation
        # For production, consider using langdetect or similar
        if re.search(r'[а-яА-Я]', content):
            return "russian"
        elif re.search(r'[一-龥]', content):
            return "chinese"
        elif re.search(r'[ぁ-んァ-ン]', content):
            return "japanese"
        else:
            return "english"
