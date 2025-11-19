"""
Document Loader Plugin
Loads and processes various document formats for RAG
"""

from typing import Dict, Any, Optional, List
import os


class DocumentLoaderPlugin:
    """Plugin for loading and processing documents"""

    name = "document_loader"
    version = "1.0.0"
    description = "Load and process documents for RAG (PDF, DOCX, TXT, MD, HTML)"
    author = "Windows AI Team"

    def __init__(self):
        self._initialized = False

    def initialize(self, config: Optional[Dict[str, Any]] = None) -> bool:
        """Initialize the Document Loader plugin"""
        try:
            self._initialized = True
            return True
        except Exception as e:
            print(f"Error initializing Document Loader plugin: {e}")
            return False

    def execute(self, action: str, params: Dict[str, Any]) -> Dict[str, Any]:
        """Execute a document loading action"""
        if not self._initialized:
            return {"success": False, "error": "Plugin not initialized"}

        try:
            if action == "load_pdf":
                return self._load_pdf(params)
            elif action == "load_docx":
                return self._load_docx(params)
            elif action == "load_text":
                return self._load_text(params)
            elif action == "load_markdown":
                return self._load_markdown(params)
            elif action == "load_html":
                return self._load_html(params)
            elif action == "load_directory":
                return self._load_directory(params)
            else:
                return {"success": False, "error": f"Unknown action: {action}"}
        except Exception as e:
            return {"success": False, "error": str(e)}

    def _load_pdf(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """Load PDF document"""
        from PyPDF2 import PdfReader

        file_path = params.get("file_path", "")

        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text()

        return {
            "success": True,
            "text": text,
            "pages": len(reader.pages),
            "source": file_path
        }

    def _load_docx(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """Load DOCX document"""
        from docx import Document

        file_path = params.get("file_path", "")

        doc = Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])

        return {
            "success": True,
            "text": text,
            "paragraphs": len(doc.paragraphs),
            "source": file_path
        }

    def _load_text(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """Load plain text file"""
        file_path = params.get("file_path", "")
        encoding = params.get("encoding", "utf-8")

        with open(file_path, "r", encoding=encoding) as f:
            text = f.read()

        return {
            "success": True,
            "text": text,
            "source": file_path
        }

    def _load_markdown(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """Load Markdown file"""
        import markdown

        file_path = params.get("file_path", "")

        with open(file_path, "r", encoding="utf-8") as f:
            md_text = f.read()

        html = markdown.markdown(md_text)

        return {
            "success": True,
            "text": md_text,
            "html": html,
            "source": file_path
        }

    def _load_html(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """Load HTML file"""
        from bs4 import BeautifulSoup

        file_path = params.get("file_path", "")

        with open(file_path, "r", encoding="utf-8") as f:
            html = f.read()

        soup = BeautifulSoup(html, "html.parser")
        text = soup.get_text()

        return {
            "success": True,
            "text": text,
            "html": html,
            "source": file_path
        }

    def _load_directory(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """Load all supported documents from directory"""
        directory = params.get("directory", "")
        recursive = params.get("recursive", False)

        documents = []
        extensions = {".pdf", ".docx", ".txt", ".md", ".html"}

        if recursive:
            for root, dirs, files in os.walk(directory):
                for file in files:
                    if any(file.endswith(ext) for ext in extensions):
                        file_path = os.path.join(root, file)
                        documents.append(file_path)
        else:
            for file in os.listdir(directory):
                if any(file.endswith(ext) for ext in extensions):
                    documents.append(os.path.join(directory, file))

        return {
            "success": True,
            "documents": documents,
            "count": len(documents)
        }

    def shutdown(self) -> bool:
        """Cleanup plugin resources"""
        self._initialized = False
        return True
